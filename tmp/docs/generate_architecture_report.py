from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/phong/Môn học/TCNS-DCD-BACK")
OUTPUT_PATH = ROOT / "output" / "doc" / "bao_cao_bai_tap_lon_kien_truc_thiet_ke_phan_mem.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "808080", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def apply_run_font(run, *, name: str = "Times New Roman", size: float = 12, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def add_paragraph(document: Document, text: str = "", *, align=None, size: float = 12, bold: bool = False, italic: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.22
    paragraph.paragraph_format.space_after = Pt(4)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        apply_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_paragraphs(document: Document, texts: list[str]) -> None:
    for text in texts:
        add_paragraph(document, text)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    size = 14 if level == 1 else 13 if level == 2 else 12
    run = paragraph.add_run(text)
    apply_run_font(run, size=size, bold=True)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        apply_run_font(run, size=11.5)


def add_number_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        apply_run_font(run, size=11.5)


def add_code_block(document: Document, text: str, *, font_size: float = 9.6) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    set_cell_borders(cell, color="AAB2BD", size="10")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    apply_run_font(run, name="Courier New", size=font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        set_cell_shading(cell, "D9EAF7")
        set_cell_borders(cell, color="7F8C8D")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(header)
        apply_run_font(run, size=11, bold=True)
        if widths:
            cell.width = Cm(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = ""
            set_cell_borders(cell, color="B0B7C3")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(value)
            apply_run_font(run, size=10.8)
            if widths:
                cell.width = Cm(widths[index])

    document.add_paragraph()


def add_placeholder_box(document: Document, title: str, note: str, *, height_cm: float = 5.2) -> None:
    p = document.add_paragraph()
    run = p.add_run(title)
    apply_run_font(run, size=11.5, bold=True)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    set_cell_shading(cell, "FBFCFC")
    set_cell_borders(cell, color="95A5A6", size="10")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(note)
    apply_run_font(run, size=11, italic=True)
    document.add_paragraph()


def add_usecase_spec(
    document: Document,
    code: str,
    name: str,
    actor: str,
    goal: str,
    preconditions: str,
    main_flow: list[str],
    alternative_flow: list[str],
    postconditions: str,
) -> None:
    add_heading(document, f"{code}. {name}", 3)
    add_table(
        document,
        ["Thuộc tính", "Nội dung"],
        [
            ["Mã use case", code],
            ["Tên use case", name],
            ["Actor chính", actor],
            ["Mục tiêu", goal],
            ["Tiền điều kiện", preconditions],
            ["Hậu điều kiện", postconditions],
        ],
        widths=[4.2, 11.8],
    )
    add_paragraph(document, "Luồng chính:", bold=True, size=11.5)
    add_number_list(document, main_flow)
    add_paragraph(document, "Luồng thay thế / ngoại lệ:", bold=True, size=11.5)
    add_bullet_list(document, alternative_flow)


def page_break(document: Document) -> None:
    document.add_page_break()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in [("Title", 17), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True

    document.core_properties.title = "Báo cáo bài tập lớn Kiến trúc và Thiết kế phần mềm"
    document.core_properties.subject = "Phân tích hệ thống Base Backend API - bản chi tiết"
    document.core_properties.author = "Codex"
    document.core_properties.language = "vi-VN"


def add_cover_page(document: Document) -> None:
    for _ in range(4):
        add_paragraph(document, "")

    add_paragraph(document, "TRƯỜNG / KHOA / BỘ MÔN", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    add_paragraph(document, "BÀI TẬP LỚN MÔN KIẾN TRÚC VÀ THIẾT KẾ PHẦN MỀM", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_paragraph(document, "Đề tài: Phân tích, mô hình hóa và đánh giá hệ thống Base Backend API", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    add_paragraph(document, "Phiên bản báo cáo chuyên sâu 30-40 trang", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)

    for _ in range(2):
        add_paragraph(document, "")

    table = document.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_rows = [
        ("Giảng viên hướng dẫn", "........................................................"),
        ("Sinh viên thực hiện", "........................................................"),
        ("Mã sinh viên", "........................................................"),
        ("Lớp", "........................................................"),
        ("Nhóm", "........................................................"),
        ("Repository khảo sát", "TCNS-DCD-BACK"),
        ("Sản phẩm khảo sát", "Base Backend API - hệ thống quản lý nội bộ TCNS"),
        ("Ngày lập báo cáo", date(2026, 4, 7).strftime("%d/%m/%Y")),
    ]
    for row_index, (left, right) in enumerate(cover_rows):
        left_cell, right_cell = table.rows[row_index].cells
        set_cell_shading(left_cell, "EAF2F8")
        set_cell_borders(left_cell, color="99A3A4")
        set_cell_borders(right_cell, color="99A3A4")
        run = left_cell.paragraphs[0].add_run(left)
        apply_run_font(run, size=11.5, bold=True)
        run = right_cell.paragraphs[0].add_run(right)
        apply_run_font(run, size=11.5)

    for _ in range(7):
        add_paragraph(document, "")

    add_paragraph(
        document,
        "Bản báo cáo này được biên soạn trực tiếp từ mã nguồn hiện có của hệ thống, bám vào cấu trúc thư mục, tuyến xử lý request, schema dữ liệu, route, controller, service, repository, database adapter và các thành phần tích hợp bên ngoài.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        italic=True,
    )


def add_toc(document: Document) -> None:
    add_heading(document, "MỤC LỤC", 1)
    add_number_list(
        document,
        [
            "I. Tổng quan và giới thiệu",
            "II. Đặc tả và mô hình hóa yêu cầu",
            "III. Cấu trúc và hành vi",
            "IV. Kiến trúc và mẫu thiết kế",
            "V. Minh chứng phần mềm",
            "Phụ lục A. Danh mục endpoint chính",
            "Phụ lục B. Tổng hợp thực thể dữ liệu",
            "Phụ lục C. Ma trận bằng chứng kiến trúc và pattern",
        ],
    )
    add_heading(document, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_table(
        document,
        ["Từ viết tắt", "Giải nghĩa"],
        [
            ["API", "Application Programming Interface"],
            ["RBAC", "Role-Based Access Control"],
            ["CRUD", "Create - Read - Update - Delete"],
            ["JWT", "JSON Web Token"],
            ["OTP", "One-Time Password"],
            ["UML", "Unified Modeling Language"],
            ["C4", "Context - Container - Component - Code"],
            ["NFR", "Non-functional Requirement"],
            ["DTO", "Data Transfer Object"],
        ],
        widths=[5, 11],
    )


def section_i(document: Document) -> None:
    add_heading(document, "I. TỔNG QUAN VÀ GIỚI THIỆU", 1)
    add_heading(document, "1. Bối cảnh bài toán", 2)
    add_paragraphs(
        document,
        [
            "Base Backend API là một hệ thống backend dùng cho bài toán quản lý nội bộ có nhiều nghiệp vụ kết hợp với nhau: xác thực và phân quyền, quản lý hồ sơ thành viên, lập lịch trực, thông báo hệ thống, quản lý thưởng - phạt, quản lý tài nguyên file và xuất báo cáo tổng hợp. Đây không phải là một API đơn chức năng, mà là một nền tảng backend mang tính tích hợp, phục vụ nhiều nhóm tác vụ quản trị trong cùng một miền nghiệp vụ.",
            "Trong môi trường vận hành nội bộ, các hệ thống dạng này thường gặp ba vấn đề lớn. Thứ nhất là dữ liệu bị phân tán, nhiều nghiệp vụ tách rời và khó đồng bộ. Thứ hai là logic phân quyền không nhất quán giữa các nhóm chức năng. Thứ ba là việc phát triển nhanh nhiều chức năng khiến hệ thống dễ rơi vào tình trạng controller quá tải, business logic lẫn với data access, từ đó gây khó khăn cho bảo trì. Codebase đang khảo sát được xây dựng để giải quyết các vấn đề đó bằng cách tổ chức thống nhất theo module nghiệp vụ nhưng vẫn duy trì một backend API chung.",
            "Từ góc nhìn sản phẩm, hệ thống hướng tới nhóm người dùng là thành viên của một tổ chức hoặc đơn vị có nhu cầu quản lý người dùng, lịch trực, truyền thông nội bộ và theo dõi hoạt động. Điều này giải thích vì sao trong cùng một hệ thống lại có sự hiện diện của các module như auth, users, duty, notifications, files, reward-penalties và reports. Mỗi module phục vụ một nhóm nhu cầu rõ ràng, nhưng dữ liệu và quyền hạn lại có liên hệ chặt chẽ với nhau.",
            "Báo cáo này có mục tiêu kép. Một mặt, tài liệu hóa đầy đủ yêu cầu và mô hình hóa phần mềm theo đúng tinh thần môn Kiến trúc và Thiết kế phần mềm. Mặt khác, báo cáo đóng vai trò là một tài liệu kỹ thuật có thể dùng lại để thuyết trình, phản biện hoặc làm cơ sở chỉnh sửa tài liệu cuối cùng của nhóm. Vì vậy, nội dung được triển khai theo hướng vừa học thuật, vừa bám sát mã nguồn đang tồn tại.",
        ],
    )

    add_heading(document, "2. Giới thiệu tổng quan về sản phẩm", 2)
    add_paragraphs(
        document,
        [
            "Sản phẩm được xây dựng bằng Node.js kết hợp Express và TypeScript. Toàn bộ hệ thống được khởi động từ một entrypoint duy nhất, khởi tạo middleware, mount route và kết nối cơ sở dữ liệu trước khi lắng nghe cổng dịch vụ. Bên trong, mã nguồn được chia theo module nghiệp vụ, mỗi module có route, controller, service và repository riêng. Mô hình tổ chức này giúp hệ thống đủ đơn giản để vận hành như một backend thống nhất, đồng thời đủ rõ ràng để tách trách nhiệm giữa các nhóm chức năng.",
            "Về chức năng, hệ thống có thể hỗ trợ trọn vẹn vòng đời làm việc của người dùng trong môi trường nội bộ. Người dùng có thể đăng ký, đăng nhập, lấy thông tin bản thân, cập nhật hồ sơ, thay đổi mật khẩu, nhận OTP khi quên mật khẩu, xem lịch trực theo tuần, đăng ký ca trực, hủy ca trực và gửi yêu cầu đổi ca. Song song, quản trị viên có thể quản lý tài khoản, thay đổi trạng thái, theo dõi thống kê, duyệt đổi ca, thao tác với dữ liệu import/export và lấy báo cáo tổng quan ở cấp hệ thống.",
            "Một điểm đáng chú ý là hệ thống không chỉ xử lý dữ liệu nghiệp vụ mà còn tích hợp dịch vụ ngoài. Cloudinary được dùng cho upload và quản lý tài nguyên media; email/SMS gateway được dùng để phát OTP khi reset mật khẩu; Swagger UI được dùng để công bố tài liệu API. Điều này cho thấy sản phẩm mang tính thực chiến, không chỉ dừng ở mức CRUD cơ bản trên cơ sở dữ liệu.",
        ],
    )

    add_heading(document, "3. Mục đích của sản phẩm", 2)
    add_bullet_list(
        document,
        [
            "Tập trung hóa các nghiệp vụ quản trị nội bộ trong một backend API thống nhất.",
            "Chuẩn hóa xác thực, phân quyền và quy tắc truy cập tài nguyên theo role/permission.",
            "Hỗ trợ quản lý vòng đời người dùng từ đăng ký, kích hoạt, cập nhật hồ sơ đến vô hiệu hóa tài khoản.",
            "Hỗ trợ vận hành lịch trực, đăng ký ca, đổi ca và thông báo cho các actor liên quan.",
            "Tạo một nền tảng backend có thể dùng làm base project cho các hệ thống cùng miền nghiệp vụ.",
        ],
    )

    add_heading(document, "4. Phạm vi và giả định khảo sát", 2)
    add_table(
        document,
        ["Nội dung", "Mô tả"],
        [
            ["Phạm vi khảo sát", "Mã nguồn backend trong repository TCNS-DCD-BACK, tập trung vào route, controller, service, repository, schema và tầng database adapter."],
            ["Ngoài phạm vi", "Frontend thực tế, cấu hình hạ tầng production, monitoring ngoài ứng dụng và các hệ thống client cụ thể."],
            ["Giả định chính", "Hệ thống được vận hành như một backend API nội bộ, người dùng tương tác qua frontend hoặc công cụ API client."],
            ["Mục tiêu tài liệu", "Vừa phục vụ báo cáo học phần vừa có thể dùng như tài liệu kỹ thuật cho chỉnh sửa tiếp theo."],
        ],
        widths=[4.2, 11.8],
    )

    add_heading(document, "5. Stakeholder và actor", 2)
    add_table(
        document,
        ["Nhóm liên quan", "Vai trò", "Kỳ vọng chính"],
        [
            ["Guest", "Người chưa đăng nhập", "Có thể đăng ký, đăng nhập và khôi phục mật khẩu nhanh, an toàn"],
            ["User", "Người dùng đã xác thực", "Quản lý hồ sơ cá nhân, lịch trực, thông báo và thao tác tệp theo quyền được cấp"],
            ["Admin", "Người quản trị hoặc nhân sự quyền cao", "Quản lý người dùng, duyệt nghiệp vụ, báo cáo và tài nguyên hệ thống"],
            ["Developer", "Nhóm phát triển backend", "Codebase dễ bảo trì, module rõ ràng, có thể mở rộng và kiểm thử"],
            ["System owner", "Người sở hữu sản phẩm", "Hệ thống ổn định, dễ triển khai, hỗ trợ đầy đủ yêu cầu vận hành"],
        ],
        widths=[3.5, 4.5, 8],
    )

    add_heading(document, "6. Các module nghiệp vụ chính", 2)
    add_table(
        document,
        ["Module", "Mục đích", "Ví dụ chức năng"],
        [
            ["auth", "Xác thực và bảo mật", "Đăng ký, đăng nhập, đổi mật khẩu, refresh token, quên mật khẩu bằng OTP"],
            ["users", "Quản lý hồ sơ người dùng", "CRUD user, cập nhật profile, quản lý trạng thái, import/export, thống kê"],
            ["duty", "Quản lý lịch trực", "Tạo ca trực, đăng ký ca, hủy ca, đổi ca, thống kê lịch trực"],
            ["notifications", "Thông báo nội bộ", "Lấy danh sách thông báo, đánh dấu đã đọc, cài đặt cấu hình nhận thông báo"],
            ["files", "Upload và quản lý tài nguyên", "Upload avatar, upload file, xem thông tin file, xóa file, thống kê storage"],
            ["reward-penalties", "Ghi nhận thưởng - phạt", "Tạo bản ghi thưởng/phạt, xem lịch sử, xem thống kê tài chính"],
            ["reports", "Báo cáo tổng quan", "Tổng hợp dữ liệu người dùng, trực, tài chính, thông báo và export báo cáo"],
        ],
        widths=[3.2, 4.8, 8],
    )

    add_heading(document, "7. Yêu cầu chức năng chi tiết", 2)
    add_table(
        document,
        ["Mã", "Yêu cầu chức năng", "Mức ưu tiên", "Actor"],
        [
            ["FR01", "Đăng ký tài khoản người dùng mới với email và mật khẩu.", "Cao", "Guest"],
            ["FR02", "Đăng nhập và nhận access token cùng refresh token.", "Cao", "Guest"],
            ["FR03", "Lấy thông tin người dùng hiện tại qua endpoint me.", "Cao", "User"],
            ["FR04", "Đổi mật khẩu khi đã đăng nhập.", "Cao", "User"],
            ["FR05", "Yêu cầu OTP để quên mật khẩu và đặt lại mật khẩu.", "Cao", "Guest / User"],
            ["FR06", "Quản trị viên tạo user mới từ giao diện quản trị hoặc API.", "Trung bình", "Admin"],
            ["FR07", "Quản trị viên cập nhật, vô hiệu hóa hoặc xóa người dùng.", "Cao", "Admin"],
            ["FR08", "Người dùng tự cập nhật hồ sơ cá nhân và avatar.", "Cao", "User"],
            ["FR09", "Hệ thống cho phép import/export dữ liệu người dùng bằng CSV/XLSX.", "Trung bình", "Admin"],
            ["FR10", "Người dùng xem lịch trực tuần hiện tại hoặc tuần được chỉ định.", "Cao", "User / Admin"],
            ["FR11", "Quản trị viên tạo và cập nhật ca trực.", "Cao", "Admin"],
            ["FR12", "Người dùng đăng ký ca trực nếu còn chỗ và không xung đột thời gian.", "Cao", "User"],
            ["FR13", "Người dùng hủy đăng ký ca trực của chính mình.", "Cao", "User"],
            ["FR14", "Người dùng tạo yêu cầu đổi ca với người dùng đích.", "Cao", "User"],
            ["FR15", "Quản trị viên duyệt hoặc từ chối yêu cầu đổi ca.", "Cao", "Admin"],
            ["FR16", "Người dùng xem, đánh dấu đã đọc, xóa thông báo và thay đổi cấu hình nhận thông báo.", "Cao", "User"],
            ["FR17", "Quản trị viên tạo bản ghi thưởng - phạt và xem thống kê tài chính.", "Trung bình", "Admin"],
            ["FR18", "Người dùng upload avatar hoặc file dùng chung; quản trị viên xem thông tin và xóa file.", "Cao", "User / Admin"],
            ["FR19", "Quản trị viên xem báo cáo tổng quan và export CSV/XLSX.", "Trung bình", "Admin"],
            ["FR20", "Hệ thống cung cấp tài liệu API và endpoint health check.", "Trung bình", "Admin / Developer"],
        ],
        widths=[1.6, 9.8, 2.8, 2.8],
    )

    add_heading(document, "8. Yêu cầu phi chức năng", 2)
    add_table(
        document,
        ["Nhóm NFR", "Nội dung chi tiết"],
        [
            ["Bảo mật", "Mật khẩu phải được băm; endpoint nhạy cảm cần xác thực JWT; phân quyền dựa trên permission; giới hạn rate cho login/register ở môi trường production."],
            ["Tính đúng đắn dữ liệu", "Schema validation cần kiểm tra ràng buộc kiểu dữ liệu, unique field, foreign key và custom validation cho password."],
            ["Tính sẵn sàng", "Ứng dụng phải có health check, xử lý lỗi tập trung và khởi tạo database trước khi nhận request."],
            ["Hiệu năng", "Các endpoint danh sách cần hỗ trợ phân trang, lọc, sort và query normalization để giảm tải cho client."],
            ["Khả năng mở rộng", "Code cần cho phép thêm module mới theo cùng cấu trúc routes/controllers/services/repositories mà không phải sửa quá nhiều thành phần lõi."],
            ["Khả năng bảo trì", "Tổ chức module rõ ràng, có shared/common, base service và base repository để giảm lặp nhưng không làm mờ trách nhiệm."],
            ["Khả năng quan sát", "Có request logging, mã lỗi nhất quán và tài liệu Swagger để hỗ trợ debug và trao đổi giữa nhóm phát triển."],
            ["Khả năng tích hợp", "Hệ thống phải giao tiếp được với Cloudinary, email/SMS gateway và các công cụ client qua giao thức HTTP chuẩn."],
        ],
        widths=[4, 12],
    )

    add_heading(document, "9. Ràng buộc và giả định kỹ thuật", 2)
    add_bullet_list(
        document,
        [
            "Hệ thống đang vận hành như một backend API duy nhất, không giả định có microservice tách riêng.",
            "Runtime hiện tại khởi tạo MongoDB thông qua database adapter, dù tài liệu môi trường có dấu vết cấu hình DB_CONNECTION cũ.",
            "Một số luồng như OTP phụ thuộc vào gateway ngoài; trong môi trường development có thể mock nếu chưa cấu hình URL thực.",
            "Phần minh chứng UI không có sẵn trong repository nên báo cáo chỉ chừa khung ảnh demo để người dùng tự chèn sau.",
        ],
    )


def section_ii(document: Document) -> None:
    add_heading(document, "II. ĐẶC TẢ VÀ MÔ HÌNH HÓA YÊU CẦU", 1)
    add_heading(document, "1. Cách tiếp cận mô hình hóa", 2)
    add_paragraphs(
        document,
        [
            "Phần này mô hình hóa yêu cầu chức năng bằng UML ở mức đủ dùng cho bài tập lớn. Do tài liệu được tạo trong định dạng DOCX để dễ chỉnh sửa, các biểu đồ được biểu diễn dưới dạng text UML thay vì ảnh tĩnh. Cách trình bày này giúp người đọc chỉnh tên actor, use case hoặc luồng xử lý trực tiếp trong tài liệu mà không phải dùng phần mềm vẽ ngoài.",
            "Các use case được lựa chọn đều là những nghiệp vụ trung tâm và có dấu vết rõ trong mã nguồn: đăng nhập, quản lý user, cập nhật profile/avatar, xem lịch trực, đăng ký ca, gửi yêu cầu đổi ca, quản lý thông báo, upload file và xuất báo cáo. Điều này giúp UML không bị chung chung mà phản ánh đúng những gì hệ thống đang hỗ trợ thực tế.",
        ],
    )

    add_heading(document, "2. Use Case Diagram tổng quát", 2)
    add_code_block(
        document,
        "+--------------------------------------------------------------------------------------+\n"
        "|                   HE THONG QUAN LY NOI BO TCNS - BACKEND API                        |\n"
        "|                                                                                      |\n"
        "| Guest  ------> (Dang ky)                                                             |\n"
        "| Guest  ------> (Dang nhap)                                                           |\n"
        "| Guest  ------> (Yeu cau OTP / Reset mat khau)                                        |\n"
        "|                                                                                      |\n"
        "| User   ------> (Xem ho so cua toi)                                                   |\n"
        "| User   ------> (Cap nhat profile / avatar)                                           |\n"
        "| User   ------> (Xem lich truc tuan)                                                  |\n"
        "| User   ------> (Dang ky ca truc)                                                     |\n"
        "| User   ------> (Huy ca truc)                                                         |\n"
        "| User   ------> (Gui yeu cau doi ca)                                                  |\n"
        "| User   ------> (Xem / danh dau / xoa thong bao)                                      |\n"
        "| User   ------> (Upload avatar / file theo quyen)                                     |\n"
        "|                                                                                      |\n"
        "| Admin  ------> (Quan ly nguoi dung)                                                  |\n"
        "| Admin  ------> (Quan ly trang thai va quyen user)                                    |\n"
        "| Admin  ------> (Import / Export du lieu user)                                        |\n"
        "| Admin  ------> (Quan ly ca truc)                                                     |\n"
        "| Admin  ------> (Duyet yeu cau doi ca)                                                |\n"
        "| Admin  ------> (Quan ly reward - penalty)                                            |\n"
        "| Admin  ------> (Xem va xuat bao cao)                                                 |\n"
        "| Admin  ------> (Quan ly tai nguyen Cloudinary)                                       |\n"
        "+--------------------------------------------------------------------------------------+",
    )

    add_heading(document, "3. Phân tích actor", 2)
    add_table(
        document,
        ["Actor", "Mô tả", "Tài nguyên thường truy cập"],
        [
            ["Guest", "Người chưa xác thực, chưa có thông tin phiên đăng nhập trong hệ thống.", "register, login, forgot-password, reset-password"],
            ["User", "Người dùng đã đăng nhập, có id, role và permission trong request context.", "me, profile, duty, notifications, upload"],
            ["Admin", "Người dùng có permission cao và có thể quản trị tài khoản, lịch trực, file và báo cáo.", "users, reports, reward-penalties, upload management, duty management"],
        ],
        widths=[3, 6.5, 6.5],
    )

    add_heading(document, "4. Đặc tả use case chi tiết", 2)
    add_usecase_spec(
        document,
        "UC01",
        "Đăng nhập hệ thống",
        "Guest",
        "Cho phép người dùng xác thực và nhận token truy cập hệ thống.",
        "Người dùng có tài khoản hợp lệ đã được tạo trong hệ thống.",
        [
            "Người dùng nhập email và mật khẩu.",
            "Client gửi POST /api/auth/login.",
            "Route kiểm tra các trường bắt buộc email và password.",
            "Controller chuyển payload cho AuthService.",
            "AuthService chuẩn hóa email và tìm user theo email.",
            "Hệ thống kiểm tra trạng thái active của tài khoản.",
            "Hệ thống so sánh mật khẩu đầu vào với mật khẩu đã băm.",
            "Nếu hợp lệ, hệ thống cập nhật lastLogin.",
            "Hệ thống sinh access token và refresh token.",
            "Controller trả response chứa thông tin user đã sanitize, quyền và token.",
        ],
        [
            "Nếu email không tồn tại, hệ thống trả lỗi 401.",
            "Nếu tài khoản bị vô hiệu hóa, hệ thống trả lỗi 401.",
            "Nếu mật khẩu sai, hệ thống trả lỗi 401 và không sinh token.",
        ],
        "Người dùng đăng nhập thành công và có thể gọi các endpoint cần xác thực.",
    )
    add_usecase_spec(
        document,
        "UC02",
        "Quản lý người dùng",
        "Admin",
        "Tạo mới, cập nhật, thay đổi trạng thái hoặc xóa người dùng.",
        "Admin đã đăng nhập và có permission phù hợp như users:create, users:update, users:delete.",
        [
            "Admin truy cập endpoint quản lý người dùng.",
            "Route kiểm tra xác thực và permission.",
            "Controller đọc request và gọi UserService hoặc service liên quan.",
            "Service kiểm tra ràng buộc dữ liệu, chuẩn hóa payload, băm mật khẩu nếu cần.",
            "Repository ghi dữ liệu vào collection users.",
            "Controller trả về thông tin user sau khi sanitize.",
        ],
        [
            "Nếu dữ liệu vi phạm schema hoặc unique field, hệ thống trả lỗi validation.",
            "Nếu admin cố thao tác hủy tài khoản của chính mình trong các luồng bị cấm, hệ thống ném ApiError tương ứng.",
            "Nếu user không tồn tại, service trả thông báo not found.",
        ],
        "Trạng thái và dữ liệu người dùng được cập nhật đúng theo action quản trị.",
    )
    add_usecase_spec(
        document,
        "UC03",
        "Cập nhật hồ sơ và avatar",
        "User",
        "Cho phép người dùng cập nhật thông tin cá nhân và thay đổi avatar.",
        "Người dùng đã đăng nhập.",
        [
            "Người dùng gửi request PUT /api/users/profile kèm dữ liệu hồ sơ và có thể kèm file avatar.",
            "Route kiểm tra xác thực và gắn middleware upload.",
            "Controller lọc ra các trường được phép cập nhật.",
            "Nếu có file, UserAvatarService xử lý upload avatar.",
            "UploadService chuẩn bị file, upload lên Cloudinary và lưu metadata file record.",
            "Service cập nhật user với avatar mới và các trường hồ sơ hợp lệ.",
            "Controller trả response user đã sanitize.",
        ],
        [
            "Nếu request không có trường nào hợp lệ và cũng không có file, hệ thống trả lỗi bad request.",
            "Nếu upload thành công nhưng ghi DB file record thất bại, service rollback xóa file đã upload trên Cloudinary.",
        ],
        "Hồ sơ người dùng và avatar được cập nhật nhất quán giữa DB và storage.",
    )
    add_usecase_spec(
        document,
        "UC04",
        "Đăng ký ca trực",
        "User",
        "Cho phép người dùng đăng ký vào một duty slot còn chỗ.",
        "Người dùng đã đăng nhập và có quyền duty:register.",
        [
            "Người dùng chọn một ca trực.",
            "Client gửi PATCH /api/duty/slots/:id/register.",
            "Route kiểm tra requireAuth và requirePermission.",
            "DutyService tìm slot theo id.",
            "Service kiểm tra slot có bị locked hay không.",
            "Service kiểm tra user đã đăng ký trước đó chưa.",
            "Service kiểm tra capacity và xung đột thời gian với các slot khác.",
            "Nếu hợp lệ, service cập nhật assignedUserIds.",
            "Service phát thông báo xác nhận đăng ký ca cho user.",
            "Controller trả slot hoặc kết quả cập nhật.",
        ],
        [
            "Nếu slot đã khóa, trả lỗi không thể đăng ký.",
            "Nếu slot đầy, trả lỗi duty slot is full.",
            "Nếu trùng khung giờ với ca khác đã nhận, trả lỗi xung đột lịch.",
        ],
        "Người dùng được gắn vào duty slot và nhận thông báo xác nhận.",
    )
    add_usecase_spec(
        document,
        "UC05",
        "Gửi yêu cầu đổi ca",
        "User",
        "Cho phép người dùng đề xuất chuyển ca trực cho một người dùng khác và chờ duyệt.",
        "Người dùng đang sở hữu ca trực và có quyền duty:update.",
        [
            "Người dùng chọn ca trực đã đăng ký và nhập người nhận ca cùng lý do.",
            "Client gửi POST /api/duty/swaps.",
            "DutyService kiểm tra slot tồn tại, requester là người đang giữ ca và target hợp lệ.",
            "Service tạo bản ghi duty swap request với trạng thái pending.",
            "Service phát thông báo cho requester, target user và nhóm approver.",
            "Controller trả bản ghi yêu cầu đổi ca vừa tạo.",
        ],
        [
            "Nếu requester không giữ ca hoặc target không hợp lệ, service trả lỗi bad request.",
            "Nếu slot/target user không tồn tại, hệ thống trả lỗi not found tương ứng.",
        ],
        "Yêu cầu đổi ca được lưu lại và chờ quản trị viên phê duyệt hoặc từ chối.",
    )
    add_usecase_spec(
        document,
        "UC06",
        "Duyệt yêu cầu đổi ca",
        "Admin",
        "Cho phép quản trị viên chấp thuận hoặc từ chối yêu cầu đổi ca.",
        "Admin đã đăng nhập và có quyền duty:approve_swap.",
        [
            "Admin xem danh sách yêu cầu đổi ca pending.",
            "Admin gửi quyết định approve hoặc reject qua PATCH /api/duty/swaps/:id/decision.",
            "DutyService phân tích payload decision và decisionNote.",
            "Nếu duyệt, service cập nhật assignedUserIds của slot và đổi trạng thái swap request.",
            "Service phát thông báo cho requester và target user.",
            "Controller trả kết quả quyết định.",
        ],
        [
            "Nếu decision không hợp lệ, service trả lỗi bad request.",
            "Nếu dữ liệu slot thay đổi khiến requester không còn sở hữu ca, service báo lỗi và không apply swap.",
        ],
        "Swap request được cập nhật trạng thái approved hoặc rejected và hệ thống đồng bộ dữ liệu ca trực tương ứng.",
    )
    add_usecase_spec(
        document,
        "UC07",
        "Quản lý thông báo",
        "User",
        "Cho phép người dùng xem, đánh dấu đã đọc, xóa thông báo và thay đổi cấu hình nhận thông báo.",
        "Người dùng đã đăng nhập.",
        [
            "Người dùng truy cập /api/notifications để lấy danh sách thông báo.",
            "NotificationService lấy dữ liệu theo userId, có hỗ trợ phân trang và đếm unread.",
            "Người dùng có thể gọi endpoint mark as read cho một notification.",
            "Người dùng có thể gọi endpoint read-all hoặc delete.",
            "Người dùng có thể cập nhật settings như shiftNotifications, approvalNotifications, emailNotifications, smsNotifications.",
        ],
        [
            "Nếu notification không thuộc về user hiện tại, service trả lỗi not found.",
            "Nếu payload settings không có khóa hợp lệ, service giữ nguyên settings cũ.",
        ],
        "Danh sách thông báo và cấu hình nhận thông báo của user được cập nhật phù hợp.",
    )
    add_usecase_spec(
        document,
        "UC08",
        "Upload file và quản lý tài nguyên",
        "User / Admin",
        "Cho phép upload avatar hoặc file chung, đồng thời cho admin quản lý tài nguyên lưu trữ.",
        "Người dùng phải xác thực; một số endpoint quản lý file yêu cầu role admin.",
        [
            "User gửi file đến endpoint upload tương ứng.",
            "Middleware upload policy kiểm tra field name, loại file và giới hạn.",
            "UploadService xử lý ảnh hoặc file chung, gọi CloudinaryStorageService để upload.",
            "FileService lưu metadata file record vào cơ sở dữ liệu.",
            "Admin có thể gọi endpoint lấy info file, xóa file hoặc dọn file cũ.",
            "Hệ thống trả thông tin tài nguyên trên Cloudinary và metadata liên quan.",
        ],
        [
            "Nếu Cloudinary chưa cấu hình, service trả lỗi cấu hình storage.",
            "Nếu ghi DB thất bại sau khi upload, hệ thống rollback xóa file đã upload.",
            "Nếu file không tồn tại, endpoint info/delete trả not found.",
        ],
        "Tài nguyên file được lưu nhất quán giữa Cloudinary và metadata store.",
    )

    add_heading(document, "5. Activity Diagram - Quy trình đăng nhập", 2)
    add_code_block(
        document,
        "Start\n"
        "  |\n"
        "Nhap email + password\n"
        "  |\n"
        "Validate request body\n"
        "  |\n"
        "Chuan hoa email\n"
        "  |\n"
        "Tim user theo email\n"
        "  |\n"
        "User ton tai?\n"
        "  |-- No --> Tra loi 401 -> End\n"
        "  |\n"
        "User active?\n"
        "  |-- No --> Tra loi 401 -> End\n"
        "  |\n"
        "Compare password\n"
        "  |\n"
        "Password dung?\n"
        "  |-- No --> Tra loi 401 -> End\n"
        "  |\n"
        "Cap nhat lastLogin\n"
        "  |\n"
        "Sinh token + refresh token\n"
        "  |\n"
        "Tra response JSON\n"
        "  |\n"
        "End",
    )
    add_paragraph(
        document,
        "Activity trên cho thấy luồng đăng nhập là một chuỗi quyết định tuyến tính với nhiều điểm fail-fast. Điều này phù hợp với yêu cầu bảo mật, vì hệ thống ngắt sớm khi thông tin xác thực không hợp lệ thay vì đi tiếp đến các bước phát token.",
    )

    add_heading(document, "6. Activity Diagram - Quy trình gửi và duyệt đổi ca", 2)
    add_code_block(
        document,
        "Start\n"
        "  |\n"
        "User chon duty slot va nhap target user + reason\n"
        "  |\n"
        "Kiem tra quyen duty:update\n"
        "  |\n"
        "Kiem tra slot ton tai va requester dang giu ca\n"
        "  |\n"
        "Kiem tra target user hop le\n"
        "  |\n"
        "Tao swap request (pending)\n"
        "  |\n"
        "Gui thong bao requester / target / approver\n"
        "  |\n"
        "Cho quyet dinh admin\n"
        "  |\n"
        "Admin approve?\n"
        "  |-- No --> Cap nhat status rejected + gui thong bao -> End\n"
        "  |\n"
        "Cap nhat assignedUserIds trong duty slot\n"
        "  |\n"
        "Cap nhat status approved + gui thong bao\n"
        "  |\n"
        "End",
    )
    add_paragraph(
        document,
        "Luồng đổi ca thể hiện một nghiệp vụ có tương tác nhiều bên: requester, target user và approver. Đây là ví dụ tốt để minh họa việc service layer điều phối nhiều repository và nhiều thông báo thay vì xử lý trực tiếp trong controller.",
    )

    add_heading(document, "7. Activity Diagram - Quy trình upload avatar", 2)
    add_code_block(
        document,
        "Start\n"
        "  |\n"
        "User gui request PUT /users/profile kem file avatar\n"
        "  |\n"
        "Middleware kiem tra auth + file policy\n"
        "  |\n"
        "Controller loc truong duoc cap nhat\n"
        "  |\n"
        "Xu ly file (prepare avatar)\n"
        "  |\n"
        "Upload len Cloudinary\n"
        "  |\n"
        "Luu file record vao DB\n"
        "  |\n"
        "Cap nhat user.avatar\n"
        "  |\n"
        "Loi khi ghi DB sau upload?\n"
        "  |-- Yes --> Rollback xoa file tren Cloudinary -> End\n"
        "  |\n"
        "Tra response user da sanitize\n"
        "  |\n"
        "End",
    )
    add_paragraph(
        document,
        "Luồng upload avatar cho thấy hệ thống đã tính đến tính nhất quán dữ liệu khi làm việc với một dịch vụ ngoài. Đây là một điểm mạnh quan trọng vì nhiều hệ thống upload file chỉ upload thành công nhưng không xử lý trường hợp ghi metadata nội bộ thất bại.",
    )


def section_iii(document: Document) -> None:
    add_heading(document, "III. CẤU TRÚC VÀ HÀNH VI", 1)
    add_heading(document, "1. Class Diagram miền dữ liệu", 2)
    add_code_block(
        document,
        "+--------------------+        1       +----------------------+\n"
        "| User               |----------------| Notification         |\n"
        "|--------------------|      0..*      |----------------------|\n"
        "| id: number         |                | id: number           |\n"
        "| name: string       |                | userId: number       |\n"
        "| email: string      |                | title: string        |\n"
        "| role: string       |                | message: string      |\n"
        "| status: string     |                | type: string         |\n"
        "| isActive: boolean  |                | isRead: boolean      |\n"
        "+--------------------+                +----------------------+\n"
        "        | 1\n"
        "        |------------------------ 1 [NotificationSetting]\n"
        "        |\n"
        "        | 1                              0..*\n"
        "        |------------------------------- [FileAsset]\n"
        "        |\n"
        "        | 0..*                        0..*\n"
        "        +----------------------------- [DutySlot]\n"
        "                                          |\n"
        "                                          | 1\n"
        "                                          |---------------- 0..* [DutySwapRequest]\n"
        "\n"
        "[RewardPenalty] * -------------------- 1 [User] : userId\n"
        "[RewardPenalty] * -------------------- 1 [User] : createdBy",
    )
    add_paragraph(
        document,
        "Class diagram miền dữ liệu tập trung vào các thực thể phản ánh dữ liệu nghiệp vụ của hệ thống. Đây là lớp mô hình hóa gần với schema hơn là với source code lớp thực thi. Cách nhìn này hữu ích khi phân tích quan hệ nghiệp vụ và xác định phạm vi tác động khi thay đổi dữ liệu.",
    )

    add_heading(document, "2. Class Diagram tầng ứng dụng", 2)
    add_code_block(
        document,
        "UserRoutes ---> UserController ---> UserService ---> UsersRepository ---> BaseRepository ---> DatabaseAdapter\n"
        "AuthRoutes ---> AuthController ---> AuthService ----> UsersRepository ----> BaseRepository ---> DatabaseAdapter\n"
        "DutyRoutes ---> DutyController ---> DutyService ----> DutySlotsRepository ---------------------> DatabaseAdapter\n"
        "UploadRoutes -> UploadController -> UploadService --> CloudinaryStorageService\n"
        "                                          |-------> FileService -------> FilesRepository ------> DatabaseAdapter\n"
        "\n"
        "BaseController <|-- UserController, AuthController, DutyController, NotificationController\n"
        "BaseService    <|-- UserService, DutyService, NotificationService, FileService, RewardPenaltyService\n"
        "BaseRepository <|-- UsersRepository, NotificationsRepository, DutySlotsRepository, FilesRepository",
    )
    add_paragraph(
        document,
        "Sơ đồ tầng ứng dụng cho thấy đây là một hệ thống phân tầng theo đúng nghĩa: route chỉ định nghĩa endpoint và middleware; controller chỉ giao tiếp HTTP; service chứa nghiệp vụ; repository truy cập dữ liệu; DatabaseAdapter che giấu chi tiết của persistence. Việc dùng BaseController, BaseService và BaseRepository giúp giảm lặp code nhưng không thay đổi trách nhiệm cốt lõi của từng layer.",
    )

    add_heading(document, "3. Mô tả các lớp/thành phần trọng tâm", 2)
    add_table(
        document,
        ["Thành phần", "Vai trò", "Nhận xét"],
        [
            ["AuthController", "Nhận request auth, validate đầu vào, gọi AuthService, trả response", "Controller mỏng, không tự làm nghiệp vụ xác thực"],
            ["AuthService", "Xử lý register, login, refresh token, change password, forgot/reset password", "Là nơi tập trung logic xác thực và quyền"],
            ["UserController", "Nhận request quản lý user, profile, avatar, stats", "Chủ yếu điều phối sang UserService hoặc UserAvatarService"],
            ["UserService", "Chứa logic nghiệp vụ về người dùng, thống kê, chuẩn hóa dữ liệu, tác động chéo sang module khác", "Là service khá lớn và giàu nghiệp vụ"],
            ["DutyService", "Quản lý duty slot, đăng ký ca, hủy ca, đổi ca, thống kê duty", "Điều phối nhiều rule nghiệp vụ và thông báo"],
            ["NotificationService", "Gửi thông báo, lấy thông báo, quản lý settings", "Gắn chặt với nhiều flow nghiệp vụ khác"],
            ["UploadService", "Điều phối upload, xử lý file, lưu metadata, rollback", "Đóng vai trò facade cho luồng upload"],
            ["BaseService", "Cung cấp CRUD chung, import/export, validation theo schema và hook before/after", "Thể hiện Template Method rõ rệt"],
            ["BaseRepository", "Chuẩn hóa data access", "Giúp service không gọi trực tiếp database"],
            ["DatabaseAdapter / MongoConnect", "Abstraction và implementation persistence", "Che giấu chi tiết của Mongoose/MongoDB"],
        ],
        widths=[3.8, 5.5, 6.7],
    )

    add_heading(document, "4. Sequence Diagram - Luồng đăng nhập", 2)
    add_code_block(
        document,
        "Guest -> AuthRoute        : POST /api/auth/login\n"
        "AuthRoute -> Validator    : validateFields(email, password)\n"
        "Validator --> AuthRoute   : hop le\n"
        "AuthRoute -> AuthController : login(req, res)\n"
        "AuthController -> AuthService : login(payload)\n"
        "AuthService -> UsersRepository : findByEmail(email)\n"
        "UsersRepository -> BaseRepository : findOne(query)\n"
        "BaseRepository -> DatabaseAdapter : findOne('users', query)\n"
        "DatabaseAdapter --> BaseRepository : user\n"
        "BaseRepository --> UsersRepository : user\n"
        "UsersRepository --> AuthService : user\n"
        "AuthService -> AuthService : comparePassword(...)\n"
        "AuthService -> UsersRepository : updateLastLogin(user.id, now)\n"
        "AuthService -> JWT helper : generateToken(), generateRefreshToken()\n"
        "JWT helper --> AuthService : token + refreshToken\n"
        "AuthService --> AuthController : authResponse\n"
        "AuthController --> Guest : 200 OK + user + permissions + token",
    )

    add_heading(document, "5. Sequence Diagram - Luồng upload avatar", 2)
    add_code_block(
        document,
        "User -> UserRoute            : PUT /api/users/profile + file\n"
        "UserRoute -> AuthMiddleware  : requireAuth\n"
        "UserRoute -> UploadMiddleware: parse avatar/image field\n"
        "UserRoute -> UserController  : updateProfile(req, res)\n"
        "UserController -> UserAvatarService : updateUserWithAvatar(...)\n"
        "UserAvatarService -> UploadService  : uploadAvatar(file, userId)\n"
        "UploadService -> FileProcessingService : prepareAvatar(file)\n"
        "UploadService -> CloudinaryStorageService : upload(buffer, folder, name)\n"
        "CloudinaryStorageService --> UploadService : cloudinary asset\n"
        "UploadService -> FileService : saveFileRecord(asset, metadata)\n"
        "FileService -> FilesRepository : create(record)\n"
        "FilesRepository -> DatabaseAdapter : create('files', record)\n"
        "UploadService --> UserAvatarService : asset + fileRecord\n"
        "UserAvatarService -> UserService/UsersRepository : update user.avatar\n"
        "UserAvatarService --> UserController : updated user\n"
        "UserController --> User : 200 OK + sanitized user",
    )

    add_heading(document, "6. Sequence Diagram - Luồng duyệt đổi ca", 2)
    add_code_block(
        document,
        "Admin -> DutyRoute         : PATCH /api/duty/swaps/:id/decision\n"
        "DutyRoute -> AuthMiddleware: requireAuth\n"
        "DutyRoute -> RBAC          : requirePermission('duty:approve_swap')\n"
        "DutyRoute -> DutyController: decideSwap(req, res)\n"
        "DutyController -> DutyService: decideSwap(id, payload, admin)\n"
        "DutyService -> DutySwapRequestsRepository : findById(id)\n"
        "DutyService -> DutyService : parse decision\n"
        "DutyService -> DutySlotsRepository : findById(dutySlotId)\n"
        "DutyService -> DutyService : applyApprovedSwapRequest() neu approve\n"
        "DutyService -> DutySlotsRepository : update assignedUserIds\n"
        "DutyService -> DutySwapRequestsRepository : update status + approvedBy + approvedAt\n"
        "DutyService -> NotificationService : notify requester / target\n"
        "DutyService --> DutyController : result\n"
        "DutyController --> Admin : 200 OK",
    )

    add_heading(document, "7. Phân tích cấu trúc và hành vi", 2)
    add_paragraphs(
        document,
        [
            "Ba sequence diagram trên thể hiện ba kiểu luồng hành vi tiêu biểu của hệ thống. Luồng đăng nhập là một luồng xác thực tuyến tính, tập trung vào kiểm tra tính hợp lệ và sinh token. Luồng upload avatar là luồng tích hợp ngoài, trong đó nhất quán dữ liệu giữa Cloudinary và database là trọng tâm. Luồng duyệt đổi ca là luồng nghiệp vụ phối hợp nhiều repository và nhiều thông báo, đại diện cho những use case phức hợp hơn mức CRUD thông thường.",
            "Từ góc nhìn thiết kế logic, hệ thống có sự phân tách khá tốt giữa phần 'orchestration' và phần 'data manipulation'. Các service như DutyService hay UploadService làm nhiệm vụ điều phối. Các repository chỉ lo query/ghi dữ liệu. Cách tách này khiến việc mô tả sequence diagram trở nên mạch lạc vì trách nhiệm của mỗi participant khá rõ. Đây là một dấu hiệu tích cực cho chất lượng thiết kế nội bộ của codebase.",
            "Một điểm nữa đáng lưu ý là các thực thể dữ liệu có nhiều quan hệ chéo xoay quanh User. Điều này làm User trở thành hạt nhân của miền dữ liệu. Trong tương lai, nếu cần refactor hoặc mở rộng hệ thống, mọi thay đổi về user lifecycle, permission hoặc status đều có khả năng lan tỏa sang các module notification, duty, files, reward-penalty và reports. Việc nhận diện mối liên kết đó ngay từ giai đoạn class diagram rất quan trọng đối với công tác bảo trì.",
        ],
    )


def section_iv(document: Document) -> None:
    add_heading(document, "IV. KIẾN TRÚC VÀ MẪU THIẾT KẾ", 1)
    add_heading(document, "1. Bằng chứng nhận diện kiến trúc hệ thống", 2)
    add_table(
        document,
        ["Bằng chứng từ code", "Ý nghĩa kiến trúc"],
        [
            ["index.ts chỉ import dotenv và src/server rồi khởi động ứng dụng", "Hệ thống có một entrypoint runtime duy nhất"],
            ["src/server.ts khởi tạo một Express app, mount middleware, mount routes và listen một cổng", "Tất cả nghiệp vụ chạy trong một process ứng dụng thống nhất"],
            ["src/routes/index.ts mount các module auth, users, upload, files, notifications, duty, reward-penalties, reports", "Các module là boundary logic trong cùng một backend, không phải service deploy riêng"],
            ["Service như UserService import trực tiếp repository và service từ module khác", "Giao tiếp nội bộ diễn ra bằng gọi hàm trực tiếp, không qua service-to-service network"],
            ["BaseRepository dùng chung DatabaseAdapter cho các module", "Tầng persistence thống nhất, không có data ownership tách theo microservice"],
        ],
        widths=[7.2, 8.8],
    )
    add_paragraph(
        document,
        "Tập hợp bằng chứng trên dẫn tới kết luận rằng hệ thống hiện tại là Monolith, cụ thể hơn là Modular Monolith. Các module nghiệp vụ được chia rõ trong mã nguồn nhưng không tách rời về runtime, deployment hay quyền sở hữu dữ liệu.",
    )

    add_heading(document, "2. So sánh Monolith, SOA và Microservice", 2)
    add_table(
        document,
        ["Tiêu chí", "Monolith", "SOA", "Microservice", "Kết quả đối với project"],
        [
            ["Số lượng ứng dụng deploy", "Một", "Nhiều service lớn", "Nhiều service nhỏ", "Một ứng dụng Express duy nhất"],
            ["Giao tiếp nội bộ", "Gọi hàm trực tiếp", "HTTP / bus / broker", "HTTP / gRPC / queue", "Import service/repository trực tiếp"],
            ["Quyền sở hữu dữ liệu", "Dùng chung hoặc tập trung", "Có thể tách", "Ưu tiên tách riêng", "Dùng chung database adapter"],
            ["Độ phức tạp vận hành", "Thấp hơn", "Trung bình", "Cao", "Phù hợp với hiện trạng codebase"],
            ["Tính độc lập phát hành", "Thấp", "Trung bình", "Cao", "Không có bằng chứng deploy độc lập từng module"],
        ],
        widths=[3.4, 2.8, 2.8, 3.2, 4.8],
    )
    add_paragraphs(
        document,
        [
            "Nếu hệ thống là SOA hoặc microservice, báo cáo phải phát hiện ít nhất một trong các dấu hiệu sau: nhiều service package chạy độc lập, nhiều entrypoint, API gateway hoặc service mesh nội bộ, giao tiếp HTTP/gRPC/queue giữa service, hoặc data store được chia theo từng service. Codebase hiện tại không cho thấy các dấu hiệu này.",
            "Ngược lại, project bộc lộ đầy đủ đặc trưng của modular monolith: cấu trúc module rõ ràng, dependency nội bộ trực tiếp, runtime thống nhất, nhưng vẫn có nhiều boundary logic để giảm độ phức tạp nhận thức. Đây là một lựa chọn phù hợp cho hệ thống có nhiều nghiệp vụ nội bộ nhưng chưa cần gánh chi phí vận hành của kiến trúc phân tán.",
        ],
    )

    add_heading(document, "3. Mẫu kiến trúc đang áp dụng", 2)
    add_table(
        document,
        ["Mẫu kiến trúc", "Có áp dụng hay không", "Phân tích"],
        [
            ["N-layered / Modular Layered", "Có", "Luồng Route -> Controller -> Service -> Repository -> Database Adapter xuất hiện rõ ở hầu hết module"],
            ["MVC", "Một phần", "Có controller và model dữ liệu nhưng không có view nội bộ; hệ thống thiên về REST API hơn là web MVC hoàn chỉnh"],
            ["CQRS", "Không", "Không tách command model, query model, handler riêng hay projection store riêng"],
            ["Clean Architecture", "Chưa đầy đủ", "Có vài abstraction và DI nhẹ nhưng chưa có dependency rule chặt và chưa tách rõ domain/application/infrastructure"],
        ],
        widths=[4.4, 3.5, 8.1],
    )
    add_code_block(
        document,
        "Request\n"
        "  -> Route\n"
        "  -> Middleware auth / validation / query\n"
        "  -> Controller\n"
        "  -> Service\n"
        "  -> Repository\n"
        "  -> DatabaseAdapter\n"
        "  -> MongoDB\n"
        "  -> Response",
        font_size=10.2,
    )
    add_paragraph(
        document,
        "Mẫu kiến trúc phù hợp nhất là Modular Layered Architecture. Mỗi module là một khối nghiệp vụ, còn bên trong mỗi module lại chia theo tầng trách nhiệm. Cách tổ chức này phù hợp với backend CRUD có thêm business rule, permission, import/export và một số tích hợp ngoài.",
    )

    add_heading(document, "4. Vì sao hệ thống chưa phải CQRS", 2)
    add_paragraphs(
        document,
        [
            "CQRS yêu cầu ít nhất một mức tách biệt rõ ràng giữa phần ghi (command) và phần đọc (query). Sự tách này có thể thể hiện qua handler riêng, model riêng, service riêng hoặc storage tối ưu khác nhau. Trong codebase hiện tại, các thao tác đọc và ghi đều đi chung qua service/repository. BaseService thậm chí cung cấp đồng thời findAll, findById, create, update, delete, search, bulk operations và import/export.",
            "Middleware parseApiQuery chỉ làm nhiệm vụ chuẩn hóa query string thành cấu trúc filter/page/limit/sort/order cho repository. Đây không phải là bằng chứng của CQRS mà chỉ là hỗ trợ query nhất quán. Không có query handler riêng, command handler riêng, event sourcing hay projection read model. Vì vậy, xếp hệ thống này vào CQRS sẽ làm sai bản chất thiết kế hiện tại.",
        ],
    )

    add_heading(document, "5. Vì sao hệ thống chưa phải Clean Architecture đầy đủ", 2)
    add_paragraphs(
        document,
        [
            "Clean Architecture yêu cầu vòng phụ thuộc hướng vào trong, tách entities, use cases, interface adapters và frameworks/infrastructure thành các lớp rõ ràng. Trong khi đó, hệ thống hiện tại vẫn theo phong cách practical layered architecture. Service import concrete repository; controller import concrete service singleton; repository import concrete db proxy. Điều này cho thấy dependency vẫn được tổ chức theo phân tầng thực dụng hơn là theo dependency rule nghiêm ngặt của Clean Architecture.",
            "Dù vậy, project có một số hạt giống tích cực. BaseRepository phụ thuộc vào interface DatabaseAdapter; một số service nhận dependency qua constructor mặc định; database được khởi tạo tập trung; nhiều quy tắc validation được gom vào service/schema thay vì rải rác. Các yếu tố này cho thấy codebase có thể tiến gần hơn tới kiến trúc sạch nếu nhu cầu bảo trì hoặc testability tăng cao trong tương lai. Nhưng tại thời điểm hiện tại, kết luận đúng hơn vẫn là 'chưa phải clean architecture đầy đủ'.",
        ],
    )

    add_heading(document, "6. Phân tích design pattern theo nhóm", 2)
    add_heading(document, "6.1. Mẫu tạo dựng", 3)
    add_table(
        document,
        ["Pattern / kỹ thuật", "Vị trí", "Lợi ích", "Giới hạn"],
        [
            ["Singleton-style module instance", "Nhiều file `export default new ...` ở controller/service/repository", "Dùng chung instance, ít boilerplate khởi tạo", "Không linh hoạt như container thật sự nếu hệ thống cần thay dependency động"],
            ["Constructor injection nhẹ", "AuthService, UploadService", "Cho phép thay dependency khi test hoặc mở rộng", "Chưa có DI container, chưa có composition root rõ ràng"],
        ],
        widths=[4.4, 4.2, 3.8, 3.6],
    )
    add_paragraph(
        document,
        "Nhóm tạo dựng trong codebase thiên về giải pháp thực dụng hơn là pattern GoF đầy đủ. Tuy nhiên, điều này phù hợp với quy mô hiện tại vì hệ thống vẫn ưu tiên sự đơn giản và khả năng đọc hiểu nhanh của developer.",
    )

    add_heading(document, "6.2. Mẫu cấu trúc", 3)
    add_table(
        document,
        ["Pattern", "Áp dụng ở đâu", "Giải thích"],
        [
            ["Repository Pattern", "BaseRepository và các repository con", "Tách service khỏi chi tiết thao tác dữ liệu và chuẩn hóa API data access"],
            ["Adapter Pattern", "DatabaseAdapter và MongoConnect", "Che giấu chi tiết persistence, giúp repository làm việc với một contract chung"],
            ["Facade-like Pattern", "UploadService", "Gói nhiều bước xử lý upload, storage, file record và rollback dưới một giao diện đơn giản"],
            ["Base class reuse", "BaseController", "Giảm lặp cho CRUD handler và chuẩn hóa cách trả response ở HTTP layer"],
        ],
        widths=[4.2, 4.6, 7.2],
    )
    add_paragraphs(
        document,
        [
            "Repository Pattern là pattern có ảnh hưởng lớn nhất ở codebase, vì nó định hình rõ quan hệ giữa service và database. Service không gọi thẳng Model của Mongoose; thay vào đó, mọi thao tác dữ liệu đều đi qua repository hoặc base repository. Điều này giúp thay đổi logic query tập trung hơn và tránh trộn lẫn persistence logic vào business layer.",
            "Adapter Pattern xuất hiện ngay dưới repository. DatabaseAdapter tạo ra một biên trừu tượng cho persistence. Dù runtime hiện tại chỉ còn MongoDB, việc có một interface chung vẫn rất có ích vì nó giảm coupling ở tầng repository và giúp giải thích kiến trúc hệ thống rõ ràng hơn trong tài liệu.",
        ],
    )

    add_heading(document, "6.3. Mẫu hành vi", 3)
    add_table(
        document,
        ["Pattern", "Áp dụng ở đâu", "Giá trị thực tế"],
        [
            ["Template Method", "BaseService", "Định nghĩa khung create/update/delete với các hook validateCreate, beforeCreate, afterCreate..."],
            ["Chain of Responsibility", "Middleware Express", "Cho phép auth, validation, query parsing và error handling chạy theo chuỗi dễ mở rộng"],
            ["Service Layer", "Hầu hết module services", "Tập trung business logic, giữ controller mỏng và repository gọn"],
        ],
        widths=[4, 5.2, 6.8],
    )
    add_paragraphs(
        document,
        [
            "Template Method trong BaseService là một ví dụ điển hình: các bước của CRUD được cố định nhưng vẫn cho phép service con chèn logic tùy chỉnh. Đây là cách thiết kế hiệu quả khi nhiều module cùng chia sẻ một khung xử lý tương tự nhưng vẫn có khác biệt nghiệp vụ.",
            "Chain of Responsibility là pattern gần như tự nhiên trong Express. Điểm mạnh của hệ thống là đã tận dụng pattern này khá sạch: route gắn requireAuth, requirePermission, validateFields hoặc validateSchema trước khi chuyển cho controller. Nhờ đó controller ít phải xử lý các nhánh kiểm tra lặp lại.",
        ],
    )

    add_heading(document, "7. C4 Model - Level 1", 2)
    add_code_block(
        document,
        "[Nguoi dung / Frontend / Quan tri vien]\n"
        "                |\n"
        "                v\n"
        "      [TCNS Base Backend API]\n"
        "          |           |           \\\n"
        "          |           |            \\__ [Email / SMS OTP Gateway]\n"
        "          |           \\_______________ [Cloudinary]\n"
        "          \\___________________________ [MongoDB]",
        font_size=10.0,
    )
    add_table(
        document,
        ["Thành phần", "Loại", "Vai trò trong bối cảnh hệ thống"],
        [
            ["Người dùng / Frontend", "Actor", "Gửi request đến backend để thao tác các nghiệp vụ quản lý nội bộ"],
            ["TCNS Base Backend API", "System", "Trung tâm xử lý xác thực, hồ sơ, lịch trực, thông báo, upload và báo cáo"],
            ["MongoDB", "External database", "Lưu dữ liệu hệ thống thông qua Mongo adapter"],
            ["Cloudinary", "External storage service", "Lưu file và ảnh đại diện, cung cấp quản lý tài nguyên media"],
            ["Email/SMS OTP Gateway", "External communication service", "Gửi OTP trong quy trình quên mật khẩu"],
        ],
        widths=[4.2, 3.5, 8.3],
    )

    add_heading(document, "8. C4 Model - Level 2", 2)
    add_code_block(
        document,
        "Container 1: Express API Monolith\n"
        "  - middleware: auth, rbac, schema validation, api query\n"
        "  - routes: auth, users, duty, notifications, files, reports, reward-penalties\n"
        "  - controllers: HTTP orchestration\n"
        "  - services: business logic\n"
        "  - repositories: data access\n"
        "  - swagger docs\n"
        "          |\n"
        "          +--> Container 2: MongoDB\n"
        "          +--> External: Cloudinary\n"
        "          +--> External: OTP Gateway",
        font_size=9.8,
    )
    add_paragraphs(
        document,
        [
            "Ở Level 2, container trung tâm duy nhất của nội bộ hệ thống là Express API Monolith. Đây là điểm rất quan trọng khi trình bày kiến trúc: routes/controllers/services/repositories không phải container tách biệt, mà là thành phần bên trong cùng một container ứng dụng.",
            "Nếu cần đi tiếp tới Level 3, khi đó mới hợp lý để bóc tách các component như AuthController, AuthService, UsersRepository, UploadService hay DutyService. Tuy nhiên, với phạm vi bài tập lớn hiện tại, Level 1 và Level 2 đã đủ để mô tả bản chất hệ thống và mối quan hệ với các external systems.",
        ],
    )

    add_heading(document, "9. Đánh giá ưu điểm, hạn chế và định hướng cải tiến", 2)
    add_table(
        document,
        ["Khía cạnh", "Ưu điểm", "Hạn chế / rủi ro", "Định hướng cải tiến"],
        [
            ["Tổ chức code", "Module hóa rõ, dễ đọc", "Service lớn có thể dần cồng kềnh", "Chia nhỏ use case trong cùng module khi cần"],
            ["Runtime", "Triển khai đơn giản", "Không có lợi ích deploy độc lập từng phần", "Chỉ tách service khi có nhu cầu thật sự"],
            ["Data access", "Repository + adapter rõ ràng", "Hiện tại runtime chỉ còn Mongo nhưng tài liệu còn dấu vết DB cũ", "Đồng bộ tài liệu môi trường với implementation"],
            ["Tái sử dụng", "BaseService và BaseRepository giảm lặp", "Có nguy cơ generic hóa quá mức", "Chỉ generic hóa phần lặp thực sự ổn định"],
            ["Tích hợp ngoài", "Cloudinary/OTP được cô lập ở service riêng", "Cần quản lý lỗi và retry kỹ hơn ở production", "Bổ sung logging/observability sâu hơn nếu đi production lớn"],
        ],
        widths=[3, 4.2, 4.6, 4.2],
    )
    add_paragraph(
        document,
        "Kết luận của phần kiến trúc là: hệ thống hiện tại phù hợp nhất với mô tả Modular Monolith kết hợp Modular Layered Architecture. Đây là một thiết kế hợp lý cho hiện trạng bài toán và cũng là nền tảng tốt để học tập về phân tầng, service layer, repository, adapter và mô hình hóa C4.",
    )


def section_v(document: Document) -> None:
    add_heading(document, "V. MINH CHỨNG PHẦN MỀM", 1)
    add_heading(document, "1. Demo sản phẩm", 2)
    add_paragraph(
        document,
        "Theo yêu cầu của người dùng, phần ảnh demo hiện tại để trống để có thể chèn thủ công sau khi chạy hệ thống. Các khung dưới đây đã được chuẩn bị sẵn với chú thích rõ ràng.",
    )
    add_placeholder_box(document, "Hình 5.1. Màn hình đăng nhập", "De trong - chen anh man hinh dang nhap sau khi chay frontend / Postman", height_cm=5.6)
    add_placeholder_box(document, "Hình 5.2. Màn hình quên mật khẩu / OTP", "De trong - chen anh man hinh forgot password hoac reset password", height_cm=5.2)
    add_placeholder_box(document, "Hình 5.3. Màn hình quản lý người dùng", "De trong - chen anh danh sach user, profile hoac action quan tri", height_cm=5.4)
    add_placeholder_box(document, "Hình 5.4. Màn hình lịch trực", "De trong - chen anh lich truc, dang ky ca hoac yeu cau doi ca", height_cm=5.4)
    add_placeholder_box(document, "Hình 5.5. Màn hình upload file", "De trong - chen anh upload avatar / file va ket qua tra ve", height_cm=5.2)
    add_placeholder_box(document, "Hình 5.6. Màn hình báo cáo tổng quan", "De trong - chen anh report overview / export CSV-XLSX", height_cm=5.4)

    add_heading(document, "2. Minh chứng mã nguồn theo nhóm chức năng", 2)
    add_table(
        document,
        ["Nhóm chức năng", "Vị trí mã nguồn", "Ý nghĩa minh chứng"],
        [
            ["Entrypoint và bootstrap", "index.ts, src/server.ts", "Chứng minh một app Express duy nhất và pipeline middleware chung"],
            ["Auth", "src/modules/auth/*", "Minh chứng đăng ký, đăng nhập, refresh token, đổi mật khẩu, OTP reset password"],
            ["Users", "src/modules/users/*", "Minh chứng CRUD người dùng, cập nhật profile, avatar, import/export, thống kê"],
            ["Duty", "src/modules/duty/*", "Minh chứng lịch trực, đăng ký/hủy ca, yêu cầu đổi ca và duyệt đổi ca"],
            ["Notifications", "src/modules/notifications/*", "Minh chứng danh sách thông báo, settings, mark as read, clear all"],
            ["Files", "src/modules/files/*", "Minh chứng upload avatar/file, metadata, storage info và cleanup Cloudinary"],
            ["Reward / Penalty", "src/modules/reward-penalties/*", "Minh chứng bản ghi thưởng - phạt và thống kê tài chính"],
            ["Reports", "src/modules/reports/*", "Minh chứng báo cáo tổng quan và export CSV/XLSX"],
            ["Shared layer", "src/shared/common/*, src/shared/repositories/*", "Minh chứng base classes, template method, repository pattern"],
            ["Persistence", "src/types/database.ts, src/database/*", "Minh chứng adapter pattern và persistence abstraction"],
        ],
        widths=[3.6, 5.6, 6.8],
    )

    add_heading(document, "3. Trích đoạn mã nguồn tiêu biểu", 2)
    add_paragraph(document, "3.1. Bootstrap server và middleware chain", bold=True, size=11.5)
    add_code_block(
        document,
        "app.use('/api', camelizeBody, logRequest, wrapJson, appendPaginationHeaders, parseApiQuery);\n"
        "setupSwagger(app);\n"
        "app.use('/api', routes);\n"
        "\n"
        "async function bootstrap() {\n"
        "  await initDatabase();\n"
        "  app.listen(PORT, () => { ... });\n"
        "}",
        font_size=9.8,
    )
    add_paragraph(
        document,
        "Đoạn mã này minh chứng rõ hệ thống có một pipeline request chung, một route root chung và một bước khởi tạo database tập trung trước khi nhận traffic.",
    )

    add_paragraph(document, "3.2. Router tổng ghép các module", bold=True, size=11.5)
    add_code_block(
        document,
        "router.use('/auth', authRoutes);\n"
        "router.use('/users', userRoutes);\n"
        "router.use('/upload', uploadRoutes);\n"
        "router.use('/files', fileRoutes);\n"
        "router.use('/notifications', notificationRoutes);\n"
        "router.use('/duty', dutyRoutes);\n"
        "router.use('/reward-penalties', rewardPenaltyRoutes);\n"
        "router.use('/reports', reportRoutes);",
        font_size=9.8,
    )
    add_paragraph(
        document,
        "Đây là bằng chứng trực tiếp của mô hình modular monolith: nhiều module nghiệp vụ được mount trong cùng một Express router.",
    )

    add_paragraph(document, "3.3. Luồng nghiệp vụ đăng nhập trong AuthService", bold=True, size=11.5)
    add_code_block(
        document,
        "const normalizedEmail = this.normalizeEmail(payload.email);\n"
        "const user = await this.userRepository.findByEmail(normalizedEmail);\n"
        "if (!user) throw ApiError.unauthorized('Invalid email or password');\n"
        "if (!user.isActive) throw ApiError.unauthorized('Account is inactive');\n"
        "const isMatch = await comparePassword(payload.password, user.password);\n"
        "await this.userRepository.updateLastLogin(user.id, loginTime);\n"
        "const token = generateToken(updatedUser.id, updatedUser.lastLogin || loginTime);",
        font_size=9.5,
    )
    add_paragraph(
        document,
        "Trích đoạn này cho thấy business logic xác thực nằm ở service layer thay vì controller, và data access được thực hiện thông qua repository.",
    )

    add_paragraph(document, "3.4. Khung CRUD chung trong BaseService", bold=True, size=11.5)
    add_code_block(
        document,
        "async create(data) {\n"
        "  const schemaValidation = await this.validateBySchema(data);\n"
        "  const customValidation = await this.validateCreate(data);\n"
        "  const transformedData = await this.beforeCreate(data);\n"
        "  const item = await this.repository.create(transformedData);\n"
        "  await this.afterCreate(item);\n"
        "  return { success: true, data: item };\n"
        "}",
        font_size=9.5,
    )
    add_paragraph(
        document,
        "Trích đoạn này là minh chứng rõ nhất cho Template Method: flow create được cố định nhưng cho phép class con chèn hook ở nhiều bước.",
    )

    add_paragraph(document, "3.5. BaseRepository và DatabaseAdapter", bold=True, size=11.5)
    add_code_block(
        document,
        "class BaseRepository {\n"
        "  constructor(collection, database = db) {\n"
        "    this.collection = collection;\n"
        "    this.database = database;\n"
        "  }\n"
        "  async findOne(query) {\n"
        "    return await this.database.findOne(this.collection, query);\n"
        "  }\n"
        "}",
        font_size=9.5,
    )
    add_paragraph(
        document,
        "Đây là minh chứng cho Repository Pattern và Adapter Pattern cùng lúc: repository chuẩn hóa truy cập dữ liệu, còn database được nhìn dưới dạng một contract trừu tượng.",
    )

    add_paragraph(document, "3.6. UploadService rollback khi lỗi ghi DB", bold=True, size=11.5)
    add_code_block(
        document,
        "const asset = await this.storageService.upload(...);\n"
        "try {\n"
        "  const fileRecord = await this.buildFileRecord(asset, prepared, options);\n"
        "  return { ...asset, fileRecord };\n"
        "} catch (error) {\n"
        "  await this.storageService.destroy(asset.publicId, asset.resourceType).catch(() => null);\n"
        "  throw error;\n"
        "}",
        font_size=9.5,
    )
    add_paragraph(
        document,
        "Đoạn này minh chứng tư duy thiết kế an toàn khi tích hợp dịch vụ ngoài: nếu phần ghi metadata thất bại thì file trên Cloudinary được rollback để tránh rác dữ liệu.",
    )

    add_heading(document, "4. Lệnh chạy, build và kiểm tra cơ bản", 2)
    add_code_block(
        document,
        "npm install\n"
        "cp .env.example .env\n"
        "npm run dev\n"
        "\n"
        "# Kiem tra app\n"
        "GET http://localhost:3000/api/health\n"
        "GET http://localhost:3000/api-docs\n"
        "\n"
        "# Build production\n"
        "npm run build\n"
        "npm start",
        font_size=10.0,
    )

    add_heading(document, "5. Ma trận liên hệ giữa báo cáo và mã nguồn", 2)
    add_table(
        document,
        ["Nội dung báo cáo", "Mã nguồn liên quan"],
        [
            ["Use case đăng nhập", "src/modules/auth/routes/auth.routes.ts, src/modules/auth/controllers/auth.controller.ts, src/modules/auth/services/auth.service.ts"],
            ["Use case quản lý user", "src/modules/users/routes/user.routes.ts, src/modules/users/controllers/user.controller.ts, src/modules/users/services/user.service.ts"],
            ["Use case lịch trực và đổi ca", "src/modules/duty/routes/duty.routes.ts, src/modules/duty/controllers/duty.controller.ts, src/modules/duty/services/duty.service.ts"],
            ["Use case thông báo", "src/modules/notifications/routes/notification.routes.ts, src/modules/notifications/services/notification.service.ts"],
            ["Use case upload", "src/modules/files/routes/upload.routes.ts, src/modules/files/services/upload.service.ts, src/modules/files/services/cloudinary-storage.service.ts"],
            ["Mẫu kiến trúc và pattern", "src/shared/common/base-service.ts, src/shared/repositories/base.repository.ts, src/types/database.ts, src/database/mongo-database.adapter.ts"],
        ],
        widths=[5.5, 10.5],
    )
    add_paragraph(
        document,
        "Phần minh chứng phần mềm cho thấy các kết luận trong báo cáo không phải suy luận cảm tính mà được rút ra từ các tuyến mã nguồn cụ thể. Đây là điểm quan trọng để bảo vệ báo cáo trước giảng viên hoặc khi nhóm cần chỉnh sửa, cập nhật tài liệu về sau.",
    )


def appendix_a(document: Document) -> None:
    add_heading(document, "PHỤ LỤC A. DANH MỤC ENDPOINT CHÍNH", 1)
    add_table(
        document,
        ["Module", "Method", "Path", "Mô tả ngắn"],
        [
            ["Auth", "POST", "/api/auth/register", "Đăng ký tài khoản"],
            ["Auth", "POST", "/api/auth/login", "Đăng nhập hệ thống"],
            ["Auth", "GET", "/api/auth/me", "Lấy thông tin người dùng hiện tại"],
            ["Auth", "PUT", "/api/auth/change-password", "Đổi mật khẩu"],
            ["Auth", "POST", "/api/auth/forgot-password", "Yêu cầu OTP reset password"],
            ["Auth", "POST", "/api/auth/reset-password", "Đặt lại mật khẩu bằng OTP"],
            ["Users", "GET", "/api/users", "Danh sách user có phân trang / filter"],
            ["Users", "POST", "/api/users", "Tạo user mới"],
            ["Users", "PUT", "/api/users/:id", "Cập nhật user"],
            ["Users", "DELETE", "/api/users/:id", "Xóa user"],
            ["Users", "PUT", "/api/users/profile", "Cập nhật profile người dùng hiện tại"],
            ["Users", "GET", "/api/users/stats/summary", "Thống kê user"],
            ["Users", "POST", "/api/users/import", "Import dữ liệu user"],
            ["Users", "GET", "/api/users/export", "Export dữ liệu user"],
            ["Duty", "GET", "/api/duty/week", "Lấy lịch trực theo tuần"],
            ["Duty", "POST", "/api/duty/slots", "Tạo ca trực"],
            ["Duty", "PATCH", "/api/duty/slots/:id/register", "Đăng ký ca trực"],
            ["Duty", "PATCH", "/api/duty/slots/:id/cancel", "Hủy ca trực"],
            ["Duty", "POST", "/api/duty/swaps", "Tạo yêu cầu đổi ca"],
            ["Duty", "PATCH", "/api/duty/swaps/:id/decision", "Duyệt / từ chối đổi ca"],
            ["Notifications", "GET", "/api/notifications", "Lấy danh sách thông báo"],
            ["Notifications", "PUT", "/api/notifications/settings", "Cập nhật cấu hình thông báo"],
            ["Files", "POST", "/api/upload/avatar", "Upload avatar"],
            ["Files", "POST", "/api/upload/general", "Upload file chung"],
            ["Files", "GET", "/api/upload/file/info", "Lấy thông tin file"],
            ["Files", "DELETE", "/api/upload/file", "Xóa file"],
            ["Reports", "GET", "/api/reports/overview", "Lấy báo cáo tổng quan"],
            ["Reports", "GET", "/api/reports/export", "Xuất báo cáo"],
            ["Reward-Penalties", "GET", "/api/reward-penalties", "Lấy lịch sử thưởng - phạt"],
            ["Reward-Penalties", "POST", "/api/reward-penalties", "Tạo bản ghi thưởng - phạt"],
        ],
        widths=[3, 2, 4.6, 6.4],
    )


def appendix_b(document: Document) -> None:
    add_heading(document, "PHỤ LỤC B. TỔNG HỢP THỰC THỂ DỮ LIỆU", 1)
    add_table(
        document,
        ["Thực thể", "Thuộc tính nổi bật", "Quan hệ chính"],
        [
            ["User", "name, email, password, role, status, avatar, isActive, lastLogin", "Liên kết với Notification, NotificationSetting, FileAsset, DutySlot, RewardPenalty"],
            ["Notification", "userId, title, message, type, category, channel, isRead", "Thuộc về một User"],
            ["NotificationSetting", "userId, shiftNotifications, approvalNotifications, emailNotifications, smsNotifications", "One-to-one với User"],
            ["FileAsset", "idFile, urlFile, uploadedBy, mimeType, provider, bytes", "uploadedBy tham chiếu User"],
            ["DutySlot", "weekStart, shiftDate, shiftLabel, startTime, endTime, capacity, assignedUserIds, status", "Nhiều User có thể gắn vào thông qua assignedUserIds"],
            ["DutySwapRequest", "dutySlotId, requesterId, targetUserId, status, decisionNote, approvedBy", "Thuộc về DutySlot và tham chiếu nhiều User theo vai trò"],
            ["RewardPenalty", "userId, type, amount, reason, eventDate, createdBy", "Tham chiếu User được thưởng/phạt và User tạo bản ghi"],
        ],
        widths=[3.4, 6.4, 6.2],
    )


def appendix_c(document: Document) -> None:
    add_heading(document, "PHỤ LỤC C. MA TRẬN BẰNG CHỨNG KIẾN TRÚC VÀ PATTERN", 1)
    add_table(
        document,
        ["Kết luận", "Bằng chứng code", "Ý nghĩa"],
        [
            ["Modular Monolith", "index.ts, src/server.ts, src/routes/index.ts", "Một runtime chung, một app Express chung, nhiều module nghiệp vụ trong cùng codebase"],
            ["Layered Architecture", "routes -> controllers -> services -> repositories", "Trách nhiệm được tách theo tầng"],
            ["Repository Pattern", "src/shared/repositories/base.repository.ts", "Chuẩn hóa data access"],
            ["Template Method", "src/shared/common/base-service.ts", "CRUD chung + hook mở rộng"],
            ["Adapter Pattern", "src/types/database.ts, src/database/mongo-database.adapter.ts", "Tách repository khỏi chi tiết DB"],
            ["Chain of Responsibility", "middleware auth, rbac, validation, parse query", "Xử lý request theo chuỗi trách nhiệm"],
            ["Facade-like UploadService", "src/modules/files/services/upload.service.ts", "Che giấu chuỗi xử lý upload phức hợp"],
        ],
        widths=[4.4, 6.4, 5.2],
    )


def count_words(document: Document) -> int:
    text_parts: list[str] = []
    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return len(" ".join(text_parts).split())


def build_document() -> Document:
    document = Document()
    configure_document(document)
    add_cover_page(document)
    page_break(document)
    add_toc(document)
    page_break(document)
    section_i(document)
    page_break(document)
    section_ii(document)
    page_break(document)
    section_iii(document)
    page_break(document)
    section_iv(document)
    page_break(document)
    section_v(document)
    page_break(document)
    appendix_a(document)
    page_break(document)
    appendix_b(document)
    page_break(document)
    appendix_c(document)
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Saved report to: {OUTPUT_PATH}")
    print(f"Approximate word count: {count_words(document)}")


if __name__ == "__main__":
    main()
